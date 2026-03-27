#!/usr/bin/env python3

import argparse
import json
import secrets
from copy import deepcopy
from pathlib import Path
from tempfile import NamedTemporaryFile
from xml.etree import ElementTree as ET
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
DRAWING_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"
OFFICE_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
COMMENT_PARTS = {
    "word/comments.xml",
    "word/commentsExtended.xml",
    "word/people.xml",
}
COMMENT_RELATIONSHIP_TYPES = {
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
    "http://schemas.microsoft.com/office/2011/relationships/commentsExtended",
    "http://schemas.microsoft.com/office/2011/relationships/people",
}
COMMENT_OVERRIDE_PARTS = {
    "/word/comments.xml",
    "/word/commentsExtended.xml",
    "/word/people.xml",
}
COMMENT_TAGS = {
    f"{{{WORD_NS}}}commentRangeStart",
    f"{{{WORD_NS}}}commentRangeEnd",
    f"{{{WORD_NS}}}commentReference",
}
PHOTO_PLACEHOLDER_DESCR = "icon.jpg"
XML_NAMESPACES = {
    "a": DRAWING_NS,
    "pic": PIC_NS,
    "r": OFFICE_REL_NS,
    "wp": WP_NS,
}


def parse_args():
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True)
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--photo")
    return parser.parse_args()


def update_para_ids(element):
    for attr in ("w14:paraId", "w14:textId"):
        key = qn(attr)
        if key in element.attrib:
            element.set(key, secrets.token_hex(4).upper())


def clone_after(anchor, template_element):
    new_element = deepcopy(template_element)
    update_para_ids(new_element)
    anchor._element.addnext(new_element)
    return Paragraph(new_element, anchor._parent)


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def text_runs(paragraph):
    return [run for run in paragraph.runs if run._r.xpath(".//w:t")]


def set_full_text(paragraph, value):
    runs = text_runs(paragraph)
    if not runs:
        return
    runs[0].text = value
    for run in runs[1:]:
        run.text = ""


def set_run_text(paragraph, index, value):
    runs = paragraph.runs
    if 0 <= index < len(runs):
        runs[index].text = value


def normalize(value):
    return (value or "").strip()


def clean_list(items):
    return [normalize(item) for item in items or [] if normalize(item)]


def chinese_index(value):
    digits = "零一二三四五六七八九"
    if value <= 10:
        return "十" if value == 10 else digits[value]
    if value < 20:
        return f"十{digits[value % 10]}"
    tens, ones = divmod(value, 10)
    prefix = f"{digits[tens]}十"
    return prefix if ones == 0 else f"{prefix}{digits[ones]}"


def join_header(period, title, role):
    parts = [normalize(period), normalize(title), normalize(role)]
    return "                        ".join(part for part in parts if part)


def format_numbered_line(index, value):
    value = normalize(value)
    return f"{index}.{value}" if value else f"{index}."


def blank_education():
    return {
        "period": "",
        "school": "",
        "major": "",
        "degree": "",
        "coreCourses": "",
        "honors": "",
        "certificates": "",
    }


def blank_experience():
    return {
        "period": "",
        "company": "",
        "role": "",
        "companySummary": "",
        "responsibilities": [],
        "achievements": [],
    }


def blank_project():
    return {
        "period": "",
        "name": "",
        "role": "",
        "summary": "",
        "highlights": [],
    }


def blank_campus():
    return {
        "mode": "bullets",
        "period": "",
        "title": "",
        "role": "",
        "bullets": [],
        "background": "",
        "responsibilities": "",
        "result": "",
    }


def keep_or_blank(items, factory):
    return items if items else [factory()]


def strip_comment_nodes(root):
    for child in list(root):
        if child.tag in COMMENT_TAGS:
            root.remove(child)
            continue

        strip_comment_nodes(child)


def ensure_content_type_for_image(root, suffix):
    ext = suffix.lstrip(".").lower()
    if ext == "jpg":
        ext = "jpeg"

    mime_type = "image/png" if ext == "png" else "image/jpeg"
    for child in root:
        if child.tag.endswith("Default") and child.attrib.get("Extension") == ext:
            return

    ET.SubElement(
        root,
        f"{{{CONTENT_TYPES_NS}}}Default",
        {"Extension": ext, "ContentType": mime_type},
    )


def find_photo_relationship_id(document_root):
    for anchor in document_root.findall(".//wp:anchor", XML_NAMESPACES):
        doc_pr = anchor.find("wp:docPr", XML_NAMESPACES)
        if doc_pr is None or doc_pr.attrib.get("descr") != PHOTO_PLACEHOLDER_DESCR:
            continue

        blip = anchor.find(".//a:blip", XML_NAMESPACES)
        if blip is None:
            continue

        return blip.attrib.get(f"{{{OFFICE_REL_NS}}}embed")

    return None


def apply_profile_photo_to_docx(path, photo_path):
    photo_path = Path(photo_path)
    if not photo_path.exists():
        return

    source_path = Path(path)
    target_name = f"word/media/profile-photo{photo_path.suffix.lower()}"
    photo_bytes = photo_path.read_bytes()

    with NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        temp_path = Path(temp_file.name)

    with ZipFile(source_path, "r") as source_zip:
        document_root = ET.fromstring(source_zip.read("word/document.xml"))
        relationship_id = find_photo_relationship_id(document_root)

        if not relationship_id:
            temp_path.unlink(missing_ok=True)
            return

        with ZipFile(temp_path, "w", compression=ZIP_DEFLATED) as target_zip:
            for item in source_zip.infolist():
                content = source_zip.read(item.filename)

                if item.filename == "word/_rels/document.xml.rels":
                    root = ET.fromstring(content)
                    for child in list(root):
                        if child.attrib.get("Id") == relationship_id:
                            child.set("Target", target_name.removeprefix("word/"))
                    content = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                elif item.filename == "[Content_Types].xml":
                    root = ET.fromstring(content)
                    ensure_content_type_for_image(root, photo_path.suffix)
                    content = ET.tostring(root, encoding="utf-8", xml_declaration=True)

                target_zip.writestr(item, content)

            target_zip.writestr(target_name, photo_bytes)

    temp_path.replace(source_path)


def scrub_comments_from_docx(path):
    source_path = Path(path)

    with NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        temp_path = Path(temp_file.name)

    with ZipFile(source_path, "r") as source_zip, ZipFile(
        temp_path,
        "w",
        compression=ZIP_DEFLATED,
    ) as target_zip:
        for item in source_zip.infolist():
            if item.filename in COMMENT_PARTS:
                continue

            content = source_zip.read(item.filename)

            if item.filename == "[Content_Types].xml":
                root = ET.fromstring(content)
                for child in list(root):
                    if child.attrib.get("PartName") in COMMENT_OVERRIDE_PARTS:
                        root.remove(child)
                content = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif item.filename == "word/_rels/document.xml.rels":
                root = ET.fromstring(content)
                for child in list(root):
                    if child.attrib.get("Type") in COMMENT_RELATIONSHIP_TYPES:
                        root.remove(child)
                content = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif item.filename.endswith(".xml") and item.filename.startswith("word/"):
                root = ET.fromstring(content)
                strip_comment_nodes(root)
                content = ET.tostring(root, encoding="utf-8", xml_declaration=True)

            target_zip.writestr(item, content)

    temp_path.replace(source_path)


def build_document(payload, template_path, output_path):
    doc = Document(str(template_path))
    paragraphs = list(doc.paragraphs)

    refs = {
        "name": paragraphs[1],
        "contact": paragraphs[3],
        "age": paragraphs[4],
        "experience_title": paragraphs[26],
        "education_anchor": paragraphs[9],
        "strength_anchor": paragraphs[18],
        "experience_anchor": paragraphs[28],
        "project_anchor": paragraphs[49],
        "campus_anchor": paragraphs[65],
    }

    prototypes = {
        "edu_header": deepcopy(paragraphs[10]._element),
        "edu_gap": deepcopy(paragraphs[11]._element),
        "edu_courses": deepcopy(paragraphs[12]._element),
        "edu_honors": deepcopy(paragraphs[13]._element),
        "edu_certs": deepcopy(paragraphs[14]._element),
        "strength_summary": deepcopy(paragraphs[19]._element),
        "strength_skill_title": deepcopy(paragraphs[20]._element),
        "strength_skill_line": deepcopy(paragraphs[21]._element),
        "strength_quality": deepcopy(paragraphs[24]._element),
        "exp_header": deepcopy(paragraphs[29]._element),
        "exp_company": deepcopy(paragraphs[30]._element),
        "exp_work_title": deepcopy(paragraphs[31]._element),
        "exp_work_line": deepcopy(paragraphs[32]._element),
        "exp_gain_title": deepcopy(paragraphs[34]._element),
        "exp_gain_line": deepcopy(paragraphs[35]._element),
        "exp_gap": deepcopy(paragraphs[45]._element),
        "proj_header": deepcopy(paragraphs[50]._element),
        "proj_summary": deepcopy(paragraphs[51]._element),
        "proj_highlight_title": deepcopy(paragraphs[53]._element),
        "proj_highlight_line": deepcopy(paragraphs[54]._element),
        "proj_gap": deepcopy(paragraphs[62]._element),
        "campus_bullet_header": deepcopy(paragraphs[66]._element),
        "campus_bullet_line": deepcopy(paragraphs[67]._element),
        "campus_detailed_header": deepcopy(paragraphs[70]._element),
        "campus_background": deepcopy(paragraphs[71]._element),
        "campus_responsibilities": deepcopy(paragraphs[72]._element),
        "campus_result": deepcopy(paragraphs[73]._element),
        "campus_gap": deepcopy(paragraphs[74]._element),
    }

    for start, end in ((66, 73), (50, 61), (29, 44), (19, 24), (10, 14)):
        for paragraph in list(doc.paragraphs)[start:end + 1]:
            remove_paragraph(paragraph)

    set_full_text(refs["name"], normalize(payload["personal"]["name"]))

    set_run_text(refs["contact"], 3, normalize(payload["personal"]["phone"]))
    set_run_text(refs["contact"], 9, normalize(payload["personal"]["email"]))

    set_run_text(refs["age"], 3, normalize(payload["personal"]["age"]))
    set_run_text(refs["age"], 9, normalize(payload["personal"]["highestEducation"]))

    set_full_text(refs["experience_title"], payload["experienceSectionTitle"])

    anchor = refs["education_anchor"]
    for education in keep_or_blank(payload["educations"], blank_education):
        header = clone_after(anchor, prototypes["edu_header"])
        period = normalize(education["period"])
        school = normalize(education["school"])
        major = normalize(education["major"])
        degree = normalize(education["degree"])
        tail = " | ".join(part for part in (major, degree) if part)
        set_full_text(
            header,
            "                      ".join(
                part for part in (period, school, tail) if part
            ),
        )
        gap = clone_after(header, prototypes["edu_gap"])
        courses = clone_after(gap, prototypes["edu_courses"])
        set_run_text(courses, 1, f"：{normalize(education['coreCourses'])}")
        honors = clone_after(courses, prototypes["edu_honors"])
        set_run_text(honors, 1, f"：{normalize(education['honors'])}")
        certs = clone_after(honors, prototypes["edu_certs"])
        set_run_text(certs, 1, f"：{normalize(education['certificates'])}")
        anchor = certs

    summary = clone_after(refs["strength_anchor"], prototypes["strength_summary"])
    set_run_text(
        summary,
        2,
        f" {normalize(payload['professionalStrengths']['summary'])}",
    )
    skill_title = clone_after(summary, prototypes["strength_skill_title"])
    anchor = skill_title
    skill_lines = payload["professionalStrengths"]["skillLines"] or [""]
    for index, item in enumerate(skill_lines, start=1):
        skill = clone_after(anchor, prototypes["strength_skill_line"])
        set_full_text(skill, f"（{index}）{normalize(item)}" if item else "")
        anchor = skill
    quality = clone_after(anchor, prototypes["strength_quality"])
    set_run_text(
        quality,
        1,
        f"：{normalize(payload['professionalStrengths']['coreQuality'])}",
    )

    anchor = refs["experience_anchor"]
    for idx, experience in enumerate(
        keep_or_blank(payload["experiences"], blank_experience),
        start=1,
    ):
        header = clone_after(anchor, prototypes["exp_header"])
        set_full_text(
            header,
            f"{chinese_index(idx)}、{join_header(experience['period'], experience['company'], experience['role'])}",
        )
        anchor = header
        company_summary = normalize(experience["companySummary"])
        if company_summary:
            company = clone_after(anchor, prototypes["exp_company"])
            set_run_text(company, 1, f"：{company_summary}")
            anchor = company

        responsibilities = clean_list(experience["responsibilities"])
        if responsibilities:
            work_title = clone_after(anchor, prototypes["exp_work_title"])
            set_run_text(work_title, 1, "：")
            anchor = work_title
        for line_index, item in enumerate(responsibilities, start=1):
            line = clone_after(anchor, prototypes["exp_work_line"])
            set_full_text(line, format_numbered_line(line_index, item))
            anchor = line

        achievements = clean_list(experience["achievements"])
        if achievements:
            gain_title = clone_after(anchor, prototypes["exp_gain_title"])
            set_run_text(gain_title, 1, "：")
            anchor = gain_title
        for line_index, item in enumerate(achievements, start=1):
            line = clone_after(anchor, prototypes["exp_gain_line"])
            set_full_text(line, format_numbered_line(line_index, item))
            anchor = line
        if idx < len(payload["experiences"] or [blank_experience()]):
            anchor = clone_after(anchor, prototypes["exp_gap"])

    anchor = refs["project_anchor"]
    projects = keep_or_blank(payload["projects"], blank_project)
    for idx, project in enumerate(projects, start=1):
        header = clone_after(anchor, prototypes["proj_header"])
        set_full_text(
            header,
            f"{chinese_index(idx)}、{join_header(project['period'], project['name'], project['role'])}",
        )
        anchor = header
        project_summary = normalize(project["summary"])
        if project_summary:
            summary = clone_after(anchor, prototypes["proj_summary"])
            set_run_text(summary, 1, f"：{project_summary}")
            anchor = summary

        highlights = clean_list(project["highlights"])
        if highlights:
            highlight_title = clone_after(anchor, prototypes["proj_highlight_title"])
            set_run_text(highlight_title, 2, "：")
            anchor = highlight_title
        for line_index, item in enumerate(highlights, start=1):
            line = clone_after(anchor, prototypes["proj_highlight_line"])
            set_full_text(line, format_numbered_line(line_index, item))
            anchor = line
        if idx < len(projects):
            anchor = clone_after(anchor, prototypes["proj_gap"])

    anchor = refs["campus_anchor"]
    campus_entries = keep_or_blank(payload["campusExperiences"], blank_campus)
    for idx, entry in enumerate(campus_entries, start=1):
        if entry["mode"] == "detailed":
            header = clone_after(anchor, prototypes["campus_detailed_header"])
            set_full_text(
                header,
                f"{chinese_index(idx)}、{join_header(entry['period'], entry['title'], entry['role'])}",
            )
            background = clone_after(header, prototypes["campus_background"])
            set_run_text(background, 1, f": {normalize(entry['background'])}")
            responsibilities = clone_after(
                background,
                prototypes["campus_responsibilities"],
            )
            set_run_text(
                responsibilities,
                1,
                f": {normalize(entry['responsibilities'])}",
            )
            result = clone_after(responsibilities, prototypes["campus_result"])
            set_run_text(result, 1, f": {normalize(entry['result'])}")
            anchor = result
        else:
            header = clone_after(anchor, prototypes["campus_bullet_header"])
            set_full_text(
                header,
                f"{chinese_index(idx)}、{join_header(entry['period'], entry['title'], entry['role'])}",
            )
            anchor = header
            bullets = entry["bullets"] or [""]
            for line_index, item in enumerate(bullets, start=1):
                bullet = clone_after(anchor, prototypes["campus_bullet_line"])
                set_full_text(bullet, format_numbered_line(line_index, item))
                anchor = bullet
        if idx < len(campus_entries):
            anchor = clone_after(anchor, prototypes["campus_gap"])

    doc.save(str(output_path))


def render_and_finalize_document(payload, template_path, output_path, photo_path=None):
    build_document(payload, template_path, output_path)

    if photo_path:
        apply_profile_photo_to_docx(output_path, photo_path)

    scrub_comments_from_docx(output_path)


def main():
    args = parse_args()
    payload = json.loads(Path(args.input).read_text(encoding="utf-8"))
    render_and_finalize_document(
        payload,
        Path(args.template),
        Path(args.output),
        args.photo,
    )


if __name__ == "__main__":
    main()
